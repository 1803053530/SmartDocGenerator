from docx import Document

# 创建新文档
doc = Document()

# 添加主标题
doc.add_heading('人工智能发展报告', level=1)

# 添加二级标题
doc.add_heading('第一章 技术背景', level=2)

# 添加三级标题
doc.add_heading('1.1 深度学习发展', level=3)

# 保存模板文件
doc.save('templates/sample.docx')

print("模板文件已生成在 templates/sample.docx")