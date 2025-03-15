import argparse
from docx import Document
import os
from openai import OpenAI
from pathlib import Path
from docx.shared import Pt, Inches

def parse_template(template_path):
    try:
        doc = Document(template_path)
        structure = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                structure.append({
                    'level': int(para.style.name.split()[-1]),
                    'text': para.text
                })
        return structure
    except Exception:
        # 返回特殊标记表示解析失败
        raise ValueError("模板结构解析失败")

def get_template_text(template_path):
    """从Word文档中提取所有文本内容"""
    doc = Document(template_path)
    return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

def generate_content(structure, theme, template_text=None):
    client = OpenAI(
        base_url="https://ark.cn-beijing.volces.com/api/v3",
        api_key="7df7f3d2-3722-401e-949d-fe6ea5c01d32"
    )
    
    # 模式切换：当有模板全文时直接处理全文
    # 统一主题参数处理
    base_prompt = f"文档主题：{theme}\n"
    
    if template_text:
        prompt = base_prompt + f"模板全文：\n{template_text}\n请根据以上主题和模板生成完整文档"
        sys_msg = "你是一位专业文档写作助手，请结合用户主题和模板内容生成文档"
    else:
        prompt = base_prompt + "文档结构：\n"
        for item in structure:
            prompt += f"{'#' * item['level']} {item['text']}\n"
        sys_msg = "你是一位专业文档写作助手，请严格按主题和结构生成内容"
    
    messages = [
        {"role": "system", "content": sys_msg},
        {"role": "user", "content": prompt}
    ]
    
    completion = client.chat.completions.create(
        model="doubao-1-5-pro-256k-250115",
        messages=messages
    )
    return completion.choices[0].message.content

def create_output_doc(template_path, content, output_path):
    try:
        doc = Document()
        
        # 直接生成完整文档模式
        sections = content.split('\n## ')
        doc.add_heading(sections[0].strip(), level=1)
        
        for section in sections[1:]:
            lines = section.split('\n')
            heading = lines[0].strip()
            body = '\n'.join(lines[1:]).strip()
            
            # 自动判断标题层级
            level = min(3, heading.count('.') + 1)  # 支持1-3级标题
            doc.add_heading(heading, level=level)
            doc.add_paragraph(body)
        
        # 设置默认样式
        style = doc.styles['Normal']
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        
        doc.save(output_path)
        print(f"文档生成成功，保存路径：{output_path}")
        
    except Exception as e:
        print(f"文档生成失败：{str(e)}")
        raise

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='智能文档生成工具')
    parser.add_argument('--template', required=True, help='模板文件路径')
    parser.add_argument('--theme', required=True, help='文档主题')
    parser.add_argument('--output', default='output.docx', help='输出文件名')
    
    args = parser.parse_args()
    
    # 确保输出目录存在
    output_dir = Path('results')
    output_dir.mkdir(exist_ok=True)
    
    try:
        # 修改文件读取方式
        template_text = get_template_text(args.template)
            
        try:
            structure = parse_template(args.template)
            content = generate_content(structure, args.theme)
        except ValueError:
            content = generate_content([], args.theme, template_text=template_text)
            
        create_output_doc(args.template, content, output_dir / args.output)
    except Exception as e:
        print(f"程序执行遇到错误：{str(e)}")
    finally:
        print("程序运行结束")